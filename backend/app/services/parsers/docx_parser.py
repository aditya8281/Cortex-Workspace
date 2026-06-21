"""Word document (.docx) parser."""

from __future__ import annotations

import logging

from backend.app.services.parsers.base import BaseParser, ParsedDocument, ParsedSection

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Extract text from Word .docx files using python-docx."""

    HEADING_MAP = {
        "Heading 1": 1,
        "Heading 2": 2,
        "Heading 3": 3,
        "Heading 4": 4,
        "Heading 5": 5,
        "Heading 6": 6,
    }

    def parse(self, file_path: str) -> ParsedDocument:
        try:
            import docx
        except ImportError:
            logger.warning("python-docx not installed — run: pip install python-docx")
            return ParsedDocument(metadata={"error": "python-docx not installed"})

        sections: list[ParsedSection] = []
        metadata: dict = {}

        try:
            document = docx.Document(file_path)

            if document.core_properties.title:
                metadata["title"] = document.core_properties.title
            if document.core_properties.author:
                metadata["author"] = document.core_properties.author

            # Extract headers
            for section in document.sections:
                for header in [section.header, section.first_page_header, section.even_page_header]:
                    if header and not header.is_linked_to_previous:
                        for para in header.paragraphs:
                            text = para.text.strip()
                            if text:
                                sections.append(
                                    ParsedSection(
                                        content=text,
                                        section_type="paragraph",
                                        title="Header",
                                    )
                                )

            # Extract body paragraphs
            for para in document.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else ""
                if style_name in self.HEADING_MAP:
                    level = self.HEADING_MAP[style_name]
                    sections.append(
                        ParsedSection(
                            title=text,
                            content=text,
                            section_type="heading",
                            level=level,
                        )
                    )
                else:
                    sections.append(
                        ParsedSection(
                            content=text,
                            section_type="paragraph",
                        )
                    )

            # Extract tables
            for table in document.tables:
                table_text = self._format_table(table)
                if table_text.strip():
                    sections.append(
                        ParsedSection(
                            content=table_text,
                            section_type="table",
                        )
                    )

            # Extract footers
            for section in document.sections:
                for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                    if footer and not footer.is_linked_to_previous:
                        for para in footer.paragraphs:
                            text = para.text.strip()
                            if text:
                                sections.append(
                                    ParsedSection(
                                        content=text,
                                        section_type="paragraph",
                                        title="Footer",
                                    )
                                )

        except Exception as e:
            logger.warning("Failed to parse DOCX %s: %s", file_path, e)
            return ParsedDocument(metadata={"error": str(e)})

        total_chars = sum(len(s.content) for s in sections)
        return ParsedDocument(sections=sections, metadata=metadata, total_chars=total_chars)

    @staticmethod
    def _format_table(table) -> str:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)

        if not rows:
            return ""

        headers = rows[0]
        data_rows = rows[1:]
        header_line = " | ".join(headers)
        separator = " | ".join(["---"] * len(headers))
        lines = [header_line, separator]
        for row in data_rows:
            lines.append(" | ".join(row))
        return "\n".join(lines)
